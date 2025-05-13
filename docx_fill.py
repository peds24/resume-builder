from docx import Document
import re
from lxml import etree
from docx.oxml.ns import qn
import docx.shared

def parse_xml(xml_string):
    """
    Custom implementation of parse_xml since it's not available in the current python-docx version
    """
    return etree.fromstring(xml_string)

def apply_bullet_formatting(doc, paragraph):
    """Apply correct formatting for bullet points (size 10) while keeping text at size 9"""
    paragraph.style = 'List Paragraph'
    
    # Create proper bullet point with size 10
    bullet_xml = parse_xml('<w:numPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                     '<w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr>')
    
    pPr = paragraph._p.get_or_add_pPr()
    
    for child in pPr.findall('.//{%s}numPr' % qn('w:')[1:-1]):
        pPr.remove(child)
    
    pPr.append(bullet_xml)
    
    # Add specific indentation for bullets
    for child in pPr.findall('.//{%s}ind' % qn('w:')[1:-1]):
        pPr.remove(child)
    
    pPr.insert(0, parse_xml('<w:ind w:left="365" w:hanging="360" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
    
    # Add number formatting for the bullets (size 10)
    if not any(child.tag == qn('w:rPr') for child in pPr):
        rPr = parse_xml('<w:rPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                        '<w:sz w:val="20"/></w:rPr>')  # 20 half-points = 10 points
        pPr.append(rPr)
    
    return paragraph

def fill_resume(template_path, output_path, data):
    doc = Document(template_path)
    
    # Helper function to strip whitespace and normalize text
    def normalize_text(text):
        return re.sub(r'\s+', ' ', text).strip()
    
    for paragraph in doc.paragraphs:
        # Normalize paragraph text to handle whitespace issues
        normalized_para_text = normalize_text(paragraph.text)
        
        # Check for exact full paragraph replacement first
        for key, value in data.items():
            normalized_key = normalize_text(key)
            
            # Case 1: Exact paragraph match for list items (qualifications)
            if normalized_para_text == normalized_key and key == '[qualifications]' and isinstance(value, list):
                p_to_replace = paragraph._element
                parent = p_to_replace.getparent()
                
                for item in value:
                    # Create an empty paragraph - important to avoid duplicate content
                    new_p = doc.add_paragraph()  
                    new_p.style = 'List Paragraph'
                    
                    # Add the text with size 9 via a run
                    run = new_p.add_run(item)
                    run.font.name = 'Times New Roman'
                    run.font.size = docx.shared.Pt(9)
                    
                    apply_bullet_formatting(doc, new_p)
                    
                    parent.insert(parent.index(p_to_replace), new_p._element)
                
                parent.remove(p_to_replace)
                break
            
            # Case 2: Jobs/projects with header and description
            elif normalized_para_text == normalized_key and isinstance(value, list) and all('header' in item and 'description' in item for item in value):
                p_to_replace = paragraph._element
                parent = p_to_replace.getparent()
                
                for item in value:
                    header = doc.add_paragraph()
                    header_run = header.add_run(item['header'])
                    # Apply Times New Roman 9pt to headers
                    header_run.font.name = 'Times New Roman'
                    header_run.font.size = docx.shared.Pt(9)
                    header_run.bold = True
                    parent.insert(parent.index(p_to_replace), header._element)
                    
                    for point in item['description']:
                        # Create an empty paragraph first (no text content)
                        bullet_p = doc.add_paragraph(style='List Paragraph')
                        
                        # Add the text with size 9 in a run
                        point_run = bullet_p.add_run(point)
                        point_run.font.name = 'Times New Roman'
                        point_run.font.size = docx.shared.Pt(9)
                        
                        # Apply bullet point formatting with size 10
                        apply_bullet_formatting(doc, bullet_p)
                        
                        parent.insert(parent.index(p_to_replace), bullet_p._element)
                
                parent.remove(p_to_replace)
                break
        
        # Case 3: Simple text replacement within runs
        for key, value in data.items():
            
            if isinstance(value, list):
                continue
            
            # Replace text in runs (handles partial paragraph matches)
            for run in paragraph.runs:
                if key in run.text:
                    run.text = run.text.replace(key, value)
                    
                    # Apply formatting if needed
                    run.font.name = 'Times New Roman'
                    run.font.size = docx.shared.Pt(9)

    doc.save(output_path)
